import os
import sys
import argparse
import requests
import json
from github import Github
from io import StringIO
import re
import tokenize
from io import BytesIO
import traceback
from docx import Document

BLACKLIST_FILENAMES = {
    '.gitignore',
    'COPYING',
    'LICENSE',
    'CHANGELOG.md',
    'CONTRIBUTORS.md',
    'CONTRIBUTING.md', 
    'FLAGS.md',
    'appveyor.yml',
    'CODEOWNERS',
    '.travis.yml',
    '.gitlab-ci.yml',
    'requirements.txt',
    'setup.py',
    'package.json',
    'package-lock.json',
    'yarn.lock',
    '.editorconfig',
    '.eslintrc',
    '.prettierrc',
    '.stylelintrc',
    'Makefile',
    'CMakeLists.txt'
}

BLACKLIST_EXTENSIONS = {
    '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp',
    '.ico', '.ttf', 'woff2', 'eot', 'woff',
    '.mp3', '.wav', '.ogg', '.flac', '.aac',
    '.mp4', '.avi', '.mov', '.wmv', '.flv', '.mkv',
    '.zip', '.rar', '.7z', '.tar', '.gz',
    '.exe', '.dll', '.so', '.dylib',
    #'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
    '.iso', '.bin', '.dat', '.img'
}

TOKEN_FILE = '.github_token.json'

def get_github_token(provided_token=None):
    if provided_token:
        # If a token is provided, save it and return it
        with open(TOKEN_FILE, 'w') as f:
            json.dump({"token": provided_token}, f)
        return provided_token
    
    # If no token is provided, try to read from file
    if os.path.exists(TOKEN_FILE):
        with open(TOKEN_FILE, 'r') as f:
            data = json.load(f)
            return data.get("token")
    
    return None

def estimate_tokens(text):
    def is_likely_code(text):
        code_patterns = [
            r'\bdef\b', r'\bclass\b', r'\bif\b.*:', r'\bfor\b.*:', r'\bwhile\b.*:',
            r'\bimport\b', r'\bfrom\b.*\bimport\b', r'^\s*@', r'^\s*#',
            r'\b(var|let|const)\b', r'\bfunction\b', r'\breturn\b',
            r'\bpublic\b', r'\bprivate\b', r'\bprotected\b', r'\bstatic\b',
            r'\b(int|float|double|string|bool|void)\b'
        ]
        return any(re.search(pattern, text, re.MULTILINE) for pattern in code_patterns)

    if is_likely_code(text):
        try:
            token_count = 0
            # Use generate_tokens instead of tokenize for more robustness
            for tok in tokenize.generate_tokens(StringIO(text).readline):
                if tok.type not in {tokenize.COMMENT, tokenize.NEWLINE, tokenize.INDENT, tokenize.DEDENT, tokenize.ENCODING}:
                    token_count += 1
            return token_count
        except tokenize.TokenError:
            # If tokenizing as Python code fails, fall back to the generic method
            pass
        except IndentationError:
            # If there's an indentation error, fall back to the generic method
            pass

    # Generic tokenization for non-code or if code tokenization fails
    words = re.findall(r'\w+|[^\w\s]', text)
    return len(words)

def get_contents_with_tokens(path, is_local=True, repo=None, github_path=""):
    contents = []
    if is_local:
        full_path = os.path.join(path, github_path)
        for item in os.listdir(full_path):
            item_path = os.path.join(full_path, item)
            relative_path = os.path.join(github_path, item)
            if os.path.isfile(item_path) and should_include_file(relative_path, item):
                try:
                    if item.lower().endswith('.docx'):
                        content = extract_text_from_docx(item_path)
                    else:
                        with open(item_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    tokens = estimate_tokens(content)
                    contents.append({
                        'path': relative_path,
                        'content': content,
                        'type': 'file',
                        'tokens': tokens
                    })
                except Exception as e:
                    print(f"Warning: Unable to process {item_path}. Error: {str(e)}")
            elif os.path.isdir(item_path):
                contents.append({
                    'path': relative_path,
                    'type': 'dir',
                    'tokens': 0
                })
    else:
        items = repo.get_contents(github_path)
        for item in items:
            if item.type == "file" and should_include_file(item.path, item.name):
                try:
                    if item.name.lower().endswith('.docx'):
                        # For GitHub repos, we need to download the file first
                        content = extract_text_from_docx(BytesIO(item.decoded_content))
                    else:
                        content = item.decoded_content.decode("utf-8")
                    tokens = estimate_tokens(content)
                    contents.append({
                        'path': item.path,
                        'content': content,
                        'type': 'file',
                        'tokens': tokens
                    })
                except Exception as e:
                    print(f"Warning: Unable to process {item.path}. Error: {str(e)}")
            elif item.type == "dir":
                contents.append({
                    'path': item.path,
                    'type': 'dir',
                    'tokens': 0
                })
    return contents

def concatenate_files_recursively(path, is_local=True, repo=None, max_tokens=None):
    total_tokens = 0
    concatenated_content = ""
    used_files = []
    to_process = [("", 0)]  # (path, level)

    while to_process:
        current_path, current_level = to_process.pop(0)
        contents = get_contents_with_tokens(path, is_local, repo, current_path)
        
        for item in contents:
            if item['type'] == 'dir':
                to_process.append((item['path'], current_level + 1))
            else:
                if max_tokens and total_tokens + item['tokens'] > max_tokens:
                    return concatenated_content, total_tokens, used_files
                
                file_path = item['path']
                file_content = item['content']
                
                concatenated_content += f"\n'''---\n{file_path}\n'''---\n{file_content}\n"
                total_tokens += item['tokens']
                used_files.append(file_path)

    return concatenated_content, total_tokens, used_files

def is_text_file(filename):
    if filename in BLACKLIST_FILENAMES:
        return False
    _, ext = os.path.splitext(filename.lower())
    return ext not in BLACKLIST_EXTENSIONS or ext == ".docx"

def should_include_file(filepath, filename):
    if filename in BLACKLIST_FILENAMES:
        return False
    if any(part.startswith('.') for part in filepath.split(os.sep)):
        return False  # Exclude hidden directories
    _, ext = os.path.splitext(filename.lower())
    return ext not in BLACKLIST_EXTENSIONS

def extract_text_from_docx(file_or_bytes):
    if isinstance(file_or_bytes, str):
        doc = Document(file_or_bytes)
    else:
        doc = Document(file_or_bytes)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def get_local_contents(path):
    for root, dirs, files in os.walk(path):
        for file in files:
            if is_text_file(file):
                yield os.path.join(root, file)

def get_local_tree(path, prefix=""):
    tree = ""
    contents = os.listdir(path)
    contents.sort()
    for item in contents:
        item_path = os.path.join(path, item)
        if os.path.isdir(item_path):
            tree += f"{prefix}├── {item}/\n"
            tree += get_local_tree(item_path, prefix + "│   ")
        else:
            tree += f"{prefix}├── {item}\n"
    return tree

def get_repo_contents(repo, path=""):
    contents = repo.get_contents(path)
    return [content for content in contents if content.type != "dir" and is_text_file(content.name)]

def get_repo_tree(repo, path="", prefix=""):
    tree = ""
    contents = repo.get_contents(path)
    for content in contents:
        if content.type == "dir":
            tree += f"{prefix}├── {content.name}/\n"
            tree += get_repo_tree(repo, content.path, prefix + "│   ")
        else:
            tree += f"{prefix}├── {content.name}\n"
    return tree

def main(path, github_token=None, token_limit=15000):
    try:
        is_local = os.path.exists(path)
        
        if not is_local:
            github_token = get_github_token(github_token)
            if not github_token:
                raise ValueError("GitHub token is required for remote repositories. Please provide a token using the -t option.")
            g = Github(github_token)
            repo = g.get_repo(path)
        else:
            repo = None

        print("Estimating repository size...")
        content, total_tokens, _ = concatenate_files_recursively(path, is_local, repo)
        
        print(f"Estimated total tokens in the repository: {total_tokens}")
        
        if total_tokens > token_limit:
            choice = input(f"The repository exceeds {token_limit} tokens. Do you want to:\n"
                           f"1. Concatenate files until reaching ~{token_limit} tokens\n"
                           "2. Convert the entire repository\n"
                           "3. Exit\n"
                           "Enter your choice (1 - 3): ")
            
            if choice == "1":
                content, actual_tokens, used_files = concatenate_files_recursively(path, is_local, repo, max_tokens=token_limit)
                output_filename = f"{'local' if is_local else repo.name}_partial_concatenated.txt"
                print(f"\nFiles included: {', '.join(used_files)}")
            elif choice == "2":
                actual_tokens = total_tokens
                output_filename = f"{'local' if is_local else repo.name}_full_concatenated.txt"
            elif choice == "3":
                print("Exiting...")
                return
            else:
                print("Invalid choice, exiting.")
                return
        else:
            actual_tokens = total_tokens
            output_filename = f"{'local' if is_local else repo.name}_concatenated.txt"

        with open(output_filename, "w", encoding="utf-8") as outfile:
            outfile.write("'''---\nRepository Structure:\n\n")
            tree = get_local_tree(path) if is_local else get_repo_tree(repo)
            #print("\n", tree, "\n")
            outfile.write(tree)
            outfile.write("\n'''---\n")
            outfile.write(content)
            outfile.write(f"\n'''---\nEstimated total tokens: {actual_tokens}\n'''---\n")

        print(f"\nConcatenation complete. Output saved to '{output_filename}'")
        print(f"Estimated total tokens: {actual_tokens}")

    except Exception as e:
        print(f"Error type: {type(e).__name__}")
        print(f"Error msg: {str(e)}")
        print("\n")
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Concatenate files from a local directory or GitHub repository.")
    parser.add_argument("path", help="Local directory path or GitHub repository full name (ex. 'owner/repo')")
    parser.add_argument("-t", "--token", help="GitHub Personal Access Token (required for GitHub repositories)")
    parser.add_argument("-l", "--limit", type=int, default=15000, help="Token limit for partial concatenation (default: 15000)")
    args = parser.parse_args()

    main(args.path, args.token, args.limit)
